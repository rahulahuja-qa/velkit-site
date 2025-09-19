from io import BytesIO
from docx import Document

def docx_from_builder(cover_letter: str, resume: dict) -> bytes:
    d = Document()
    d.add_heading("Cover Letter", level=1)
    d.add_paragraph(cover_letter or "")
    d.add_page_break()
    d.add_heading("Resume", level=1)
    if resume:
        if resume.get("summary"):
            d.add_heading("Summary", level=2)
            d.add_paragraph(resume["summary"])
        if resume.get("experience"):
            d.add_heading("Experience", level=2)
            for e in resume["experience"]:
                title = f'{e.get("title","")} – {e.get("company","")}'.strip(" –")
                d.add_paragraph(title, style="List Bullet")
                for b in e.get("bullets") or []:
                    d.add_paragraph(b, style="List Number")
        if resume.get("skills"):
            d.add_heading("Skills", level=2)
            d.add_paragraph(", ".join(resume["skills"]))
    buf = BytesIO(); d.save(buf); return buf.getvalue()

def docx_from_review(payload: dict) -> bytes:
    d = Document()
    d.add_heading("Resume Review", level=1)
    d.add_paragraph(f"ATS Score: {payload.get('ats_score','-')}")
    miss = payload.get("missing_keywords") or []
    d.add_heading("Missing Keywords", level=2)
    if miss:
        for k in miss: d.add_paragraph(k, style="List Bullet")
    else:
        d.add_paragraph("None detected.")
    fb = payload.get("section_feedback") or {}
    d.add_heading("Section Feedback", level=2)
    for k, v in fb.items():
        d.add_paragraph(f"{k}: {v}")
    rew = payload.get("rewrite_suggestions") or []
    d.add_heading("Rewrite Suggestions", level=2)
    for r in rew: d.add_paragraph(r, style="List Number")
    buf = BytesIO(); d.save(buf); return buf.getvalue()

def docx_from_trainings(plan: list) -> bytes:
    d = Document()
    d.add_heading("Training & Certification Plan", level=1)
    for item in plan or []:
        d.add_heading(item.get("title",""), level=2)
        meta = f'{item.get("provider","")} • {item.get("level","")} • ~{item.get("estimated_hours","?")}h'
        d.add_paragraph(meta)
        d.add_paragraph(item.get("why",""))
    buf = BytesIO(); d.save(buf); return buf.getvalue()
